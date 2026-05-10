﻿import json
import asyncio
import logging
import mimetypes
import base64
import hashlib
import hmac
import time
from pathlib import Path
from io import BytesIO
from datetime import datetime, timedelta, timezone
from email.utils import formatdate
from typing import Optional
from urllib.parse import quote, urlencode
from uuid import uuid4

from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, Depends, Header, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles
import httpx
from jose import JWTError, jwt
from passlib.context import CryptContext
from pydantic import BaseModel, Field
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.orm import Session

from config import (
    SECRET_KEY,
    ALGORITHM,
    ACCESS_TOKEN_EXPIRE_MINUTES,
    DEEPSEEK_API_KEY,
    DEEPSEEK_MODEL,
    REBUILD_KNOWLEDGE_INDEX_ON_STARTUP,
    OSS_ACCESS_KEY_ID,
    OSS_ACCESS_KEY_SECRET,
    OSS_BUCKET,
    OSS_ENDPOINT,
    VISION_API_KEY,
    VISION_BASE_URL,
    VISION_MODEL,
    VISION_OSS_URL_EXPIRES_SECONDS,
    TEXT_FALLBACK_ENABLED,
    TEXT_FALLBACK_API_KEY,
    TEXT_FALLBACK_BASE_URL,
    TEXT_FALLBACK_MODEL,
    MEMORY_RECENT_MAX_CHARS,
    MEMORY_SUMMARY_MAX_CHARS,
    MEMORY_SUMMARY_TRIGGER_TURNS,
    MEMORY_WINDOW_TURNS,
)
from database import get_db, init_db, SessionLocal
from models import User, Conversation, Message, KnowledgeFile, KnowledgeBase, ChatTraceSession, _new_id
from retrieval import deepseek_chat_url, normalize_deepseek_model, retrieve_knowledge
from ragas_eval import schedule_ragas_evaluation
from checkpointer import delete_thread_checkpoints, list_threads
from learning_trace import (
    TraceRecorder,
    append_trace_event,
    compact_trace_reference,
    get_trace_snapshot,
    serialize_trace_session as _serialize_trace_session,
    summarize_messages,
    summarize_text,
)


logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# lifespan
# ---------------------------------------------------------------------------
@asynccontextmanager
async def lifespan(_app: FastAPI):
    init_db()
    _seed_default_users()
    if REBUILD_KNOWLEDGE_INDEX_ON_STARTUP:
        _rebuild_existing_knowledge_index()
    yield


app = FastAPI(title="RAG API", lifespan=lifespan)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
UPLOAD_DIR = Path(__file__).resolve().parent / "uploads"
AVATAR_DIR = UPLOAD_DIR / "avatars"
AVATAR_DIR.mkdir(parents=True, exist_ok=True)
app.mount("/uploads", StaticFiles(directory=str(UPLOAD_DIR)), name="uploads")


@app.get("/")
def root():
    return {"status": "ok", "service": "RAG API"}


@app.get("/health")
def health():
    return {"status": "ok"}


def _seed_default_users():
    db = SessionLocal()
    try:
        if not db.query(User).first():
            for username, password in [("admin", "admin123"), ("demo", "demo123")]:
                db.add(User(
                    username=username,
                    password_hash=pwd_context.hash(password),
                ))
            db.commit()
    finally:
        db.close()


def _get_default_knowledge_base(db: Session) -> KnowledgeBase:
    knowledge_base = db.query(KnowledgeBase).order_by(KnowledgeBase.id.asc()).first()
    if knowledge_base:
        return knowledge_base
    knowledge_base = KnowledgeBase(name="默认知识库")
    db.add(knowledge_base)
    db.commit()
    db.refresh(knowledge_base)
    return knowledge_base


def _resolve_knowledge_base(db: Session, knowledge_base_id: Optional[int]) -> KnowledgeBase:
    if knowledge_base_id:
        knowledge_base = db.query(KnowledgeBase).filter_by(id=knowledge_base_id).first()
        if knowledge_base:
            return knowledge_base
        raise HTTPException(404, "知识库不存在")
    return _get_default_knowledge_base(db)


def _rebuild_existing_knowledge_index():
    db = SessionLocal()
    try:
        from chroma_client import add_chunks, delete_file_chunks

        for entry in db.query(KnowledgeFile).filter(KnowledgeFile.knowledge_base_id.isnot(None)).all():
            delete_file_chunks(entry.id)
            chunks = _chunk_text(entry.content or "", entry.id)
            add_chunks(chunks, entry.id, entry.name, entry.knowledge_base_id)
    finally:
        db.close()


# ---------------------------------------------------------------------------
# auth helpers
# ---------------------------------------------------------------------------

def create_token(username: str) -> str:
    payload = {
        "sub": username,
        "exp": datetime.now(timezone.utc) + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES),
    }
    return jwt.encode(payload, SECRET_KEY, algorithm=ALGORITHM)


def _decode_token(authorization: str) -> str:
    """Extract and validate the token, return username."""
    if not authorization:
        raise HTTPException(401, "Missing authorization header")
    scheme, _, token = authorization.partition(" ")
    if scheme.lower() != "bearer" or not token:
        raise HTTPException(401, "Invalid authorization header")
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        return payload.get("sub", "")
    except JWTError:
        raise HTTPException(401, "Invalid token")


def get_current_user(authorization: str = Header(""), db: Session = Depends(get_db)) -> User:
    username = _decode_token(authorization)
    user = db.query(User).filter_by(username=username).first()
    if not user:
        raise HTTPException(401, "User not found")
    return user


# ---------------------------------------------------------------------------
# schemas
# ---------------------------------------------------------------------------

class LoginRequest(BaseModel):
    username: str
    password: str


class LoginResponse(BaseModel):
    token: str
    username: str


class PasswordUpdate(BaseModel):
    old_password: str
    new_password: str


class ChatRequest(BaseModel):
    conversation_id: Optional[str] = None
    knowledge_base_id: Optional[int] = None
    question: str
    attachments: list[dict] = Field(default_factory=list)


class KnowledgeBaseRequest(BaseModel):
    name: str


# ---------------------------------------------------------------------------
# auth endpoints
# ---------------------------------------------------------------------------


@app.post("/api/auth/login")
def login(body: LoginRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter_by(username=body.username).first()
    if not user or not pwd_context.verify(body.password, user.password_hash):
        raise HTTPException(401, "用户名或密码错误")
    token = create_token(body.username)
    return LoginResponse(token=token, username=body.username)


@app.post("/api/auth/logout")
def logout(_user: User = Depends(get_current_user)):
    return {"message": "ok"}


# ---------------------------------------------------------------------------
# user endpoints
# ---------------------------------------------------------------------------


@app.get("/api/user/profile")
def get_profile(user: User = Depends(get_current_user)):
    return {
        "username": user.username,
        "avatar": user.avatar or "",
        "created_at": user.created_at.isoformat() if user.created_at else "",
    }


@app.put("/api/user/password")
def update_password(body: PasswordUpdate, user: User = Depends(get_current_user),
                    db: Session = Depends(get_db)):
    if not pwd_context.verify(body.old_password, user.password_hash):
        raise HTTPException(400, "当前密码不正确")
    user.password_hash = pwd_context.hash(body.new_password)
    db.commit()
    return {"message": "密码修改成功"}


@app.post("/api/user/avatar")
async def upload_avatar(file: UploadFile = File(...),
                        user: User = Depends(get_current_user),
                        db: Session = Depends(get_db)):
    allowed_types = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/webp": ".webp",
    }
    content_type = file.content_type or ""
    ext = allowed_types.get(content_type)
    if not ext:
        raise HTTPException(400, "仅支持 png、jpg、jpeg、webp 格式头像")

    content = await file.read()
    if len(content) > 2 * 1024 * 1024:
        raise HTTPException(400, "头像文件不能超过 2MB")

    filename = f"user_{user.id}_{int(datetime.now().timestamp())}{ext}"
    target = AVATAR_DIR / filename
    target.write_bytes(content)

    user.avatar = f"/uploads/avatars/{filename}"
    db.commit()
    return {"avatar": user.avatar}


# ---------------------------------------------------------------------------
# chat endpoints
# ---------------------------------------------------------------------------


@app.get("/api/chat/conversations")
def list_conversations(user: User = Depends(get_current_user),
                       db: Session = Depends(get_db)):
    rows = (
        db.query(Conversation)
        .filter_by(user_id=user.id)
        .order_by(Conversation.updated_at.desc())
        .all()
    )
    return [
        {
            "id": c.id,
            "title": c.title,
            "knowledge_base_id": c.knowledge_base_id,
            "knowledge_base_name": c.knowledge_base.name if c.knowledge_base else "",
            "created_at": c.created_at.isoformat() if c.created_at else "",
            "updated_at": c.updated_at.isoformat() if c.updated_at else "",
        }
        for c in rows
    ]


@app.get("/api/chat/conversations/{cid}")
def get_messages(cid: str, user: User = Depends(get_current_user),
                 db: Session = Depends(get_db)):
    conv = db.query(Conversation).filter_by(id=cid, user_id=user.id).first()
    if not conv:
        raise HTTPException(404, "对话不存在")
    return [_serialize_message(m) for m in conv.messages]


@app.get("/api/chat/traces/{trace_id}")
def get_chat_trace(trace_id: str, user: User = Depends(get_current_user)):
    trace = get_trace_snapshot(trace_id, user_id=user.id)
    if not trace:
        raise HTTPException(404, "Trace 不存在")
    return trace


@app.get("/api/chat/messages/{message_id}/trace")
def get_message_trace(message_id: int, user: User = Depends(get_current_user),
                      db: Session = Depends(get_db)):
    message = (
        db.query(Message)
        .join(Conversation, Message.conversation_id == Conversation.id)
        .filter(Message.id == message_id, Conversation.user_id == user.id)
        .first()
    )
    if not message:
        raise HTTPException(404, "消息不存在")

    retrieval_trace = _loads_json(message.retrieval_trace, {})
    learning_trace = retrieval_trace.get("learning_trace", {}) if isinstance(retrieval_trace, dict) else {}
    trace_id = learning_trace.get("trace_id")
    if trace_id:
        trace = get_trace_snapshot(trace_id, user_id=user.id)
        if trace:
            return trace

    trace_session = db.query(ChatTraceSession).filter_by(message_id=message_id, user_id=user.id).first()
    if trace_session:
        return _serialize_trace_session(trace_session)
    if learning_trace:
        return learning_trace
    raise HTTPException(404, "该消息暂无流程 Trace")


@app.delete("/api/chat/conversations/{cid}")
def delete_conversation(cid: str, user: User = Depends(get_current_user),
                        db: Session = Depends(get_db)):
    conv = db.query(Conversation).filter_by(id=cid, user_id=user.id).first()
    if not conv:
        raise HTTPException(404, "对话不存在")
    db.delete(conv)
    db.commit()
    # clean checkpointer state
    delete_thread_checkpoints(cid)
    return {"message": "ok"}


class RenameRequest(BaseModel):
    title: str


@app.put("/api/chat/conversations/{cid}")
def rename_conversation(cid: str, body: RenameRequest, user: User = Depends(get_current_user),
                        db: Session = Depends(get_db)):
    conv = db.query(Conversation).filter_by(id=cid, user_id=user.id).first()
    if not conv:
        raise HTTPException(404, "对话不存在")
    conv.title = body.title
    db.commit()
    return {"message": "ok"}


def _serialize_knowledge_base(item: KnowledgeBase) -> dict:
    return {
        "id": item.id,
        "name": item.name,
        "file_count": len(item.files),
        "created_at": item.created_at.isoformat() if item.created_at else "",
        "updated_at": item.updated_at.isoformat() if item.updated_at else "",
    }


@app.get("/api/knowledge-bases")
def list_knowledge_bases(_user: User = Depends(get_current_user),
                         db: Session = Depends(get_db)):
    rows = db.query(KnowledgeBase).order_by(KnowledgeBase.created_at.asc()).all()
    return [_serialize_knowledge_base(item) for item in rows]


@app.post("/api/knowledge-bases")
def create_knowledge_base(body: KnowledgeBaseRequest,
                          _user: User = Depends(get_current_user),
                          db: Session = Depends(get_db)):
    name = body.name.strip()
    if not name:
        raise HTTPException(400, "知识库名称不能为空")
    if db.query(KnowledgeBase).filter_by(name=name).first():
        raise HTTPException(400, "知识库名称已存在")
    entry = KnowledgeBase(name=name)
    db.add(entry)
    db.commit()
    db.refresh(entry)
    return _serialize_knowledge_base(entry)


@app.put("/api/knowledge-bases/{kid}")
def rename_knowledge_base(kid: int, body: KnowledgeBaseRequest,
                          _user: User = Depends(get_current_user),
                          db: Session = Depends(get_db)):
    entry = db.query(KnowledgeBase).filter_by(id=kid).first()
    if not entry:
        raise HTTPException(404, "知识库不存在")
    name = body.name.strip()
    if not name:
        raise HTTPException(400, "知识库名称不能为空")
    duplicated = db.query(KnowledgeBase).filter(KnowledgeBase.name == name, KnowledgeBase.id != kid).first()
    if duplicated:
        raise HTTPException(400, "知识库名称已存在")
    entry.name = name
    db.commit()
    db.refresh(entry)
    return _serialize_knowledge_base(entry)


@app.delete("/api/knowledge-bases/{kid}")
def delete_knowledge_base(kid: int, _user: User = Depends(get_current_user),
                          db: Session = Depends(get_db)):
    entry = db.query(KnowledgeBase).filter_by(id=kid).first()
    if not entry:
        raise HTTPException(404, "知识库不存在")
    if db.query(KnowledgeBase).count() <= 1:
        raise HTTPException(400, "至少保留一个知识库")

    target = db.query(KnowledgeBase).filter(KnowledgeBase.id != kid).order_by(KnowledgeBase.id.asc()).first()
    for conversation in db.query(Conversation).filter_by(knowledge_base_id=kid).all():
        conversation.knowledge_base_id = target.id
    for file_entry in db.query(KnowledgeFile).filter_by(knowledge_base_id=kid).all():
        from chroma_client import delete_file_chunks
        delete_file_chunks(file_entry.id)
        db.delete(file_entry)
    db.delete(entry)
    db.commit()
    return {"message": "ok", "fallback_knowledge_base_id": target.id}


@app.post("/api/chat/attachments")
async def upload_chat_attachment(file: UploadFile = File(...),
                                 _user: User = Depends(get_current_user)):
    allowed_types = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/webp": ".webp",
    }
    content_type = file.content_type or mimetypes.guess_type(file.filename or "")[0] or ""
    ext = allowed_types.get(content_type)
    if not ext:
        raise HTTPException(400, "仅支持 png、jpg、jpeg、webp 图片")

    content = await file.read()
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(400, "图片不能超过 5MB")

    object_key = f"rag-chat/{datetime.now().strftime('%Y/%m/%d')}/{uuid4().hex}{ext}"
    try:
        await _put_oss_object(object_key, content, content_type)
    except Exception as exc:
        raise HTTPException(500, f"OSS 上传失败：{exc}")

    return {
        "name": file.filename or f"image{ext}",
        "size": len(content),
        "content_type": content_type,
        "object_key": object_key,
        "url": _public_oss_url(object_key),
    }


@app.post("/api/chat/stream")
async def stream_chat(body: ChatRequest, authorization: str = Header("")):
    username = _decode_token(authorization)
    db = SessionLocal()
    user = db.query(User).filter_by(username=username).first()
    if not user:
        db.close()
        raise HTTPException(401, "User not found")
    trace = TraceRecorder(user_id=user.id)
    try:
        trace.add(
            "request_received",
            "stream_chat",
            creates={"trace_id": trace.trace_id},
            params={
                "conversation_id": body.conversation_id,
                "knowledge_base_id": body.knowledge_base_id,
                "question": body.question,
                "attachments_count": len(body.attachments or []),
            },
            result={"username": username},
            note="后端收到一次聊天请求，先建立 trace_id，后续所有步骤都会挂到这次请求下面。",
        )
        cid = body.conversation_id
        knowledge_base = _resolve_knowledge_base(db, body.knowledge_base_id)
        raw_question = (body.question or "").strip()
        display_question = raw_question or ("请分析这张图片" if body.attachments else "")
        if not display_question:
            trace.add(
                "request_rejected",
                "stream_chat",
                uses={"raw_question": raw_question, "attachments_count": len(body.attachments or [])},
                result={"error": "问题不能为空"},
                note="没有文字问题，也没有图片附件，无法继续进入 RAG 流程。",
            )
            trace.finish("failed")
            raise HTTPException(400, "问题不能为空")
        trace.add(
            "input_normalized",
            "stream_chat",
            creates={"raw_question": raw_question, "display_question": display_question},
            result={"knowledge_base_id": knowledge_base.id, "knowledge_base_name": knowledge_base.name},
            note="系统整理用户输入，并确定本次请求要使用哪个知识库。",
        )

        effective_question, image_analysis = await _build_effective_question(
            raw_question,
            body.attachments,
        )
        trace.add(
            "effective_question_built",
            "_build_effective_question",
            params={"raw_question": raw_question, "attachments_count": len(body.attachments or [])},
            creates={
                "effective_question": effective_question,
                "image_analysis_status": image_analysis.get("status", ""),
                "image_description": image_analysis.get("description", ""),
            },
            result={"image_analysis_error": image_analysis.get("error", "")},
            note="如果有图片，系统会先把图片转成文字描述，再与用户问题合并为真正用于检索和生成的问题。",
        )
        if body.attachments and image_analysis.get("status") == "failed" and not raw_question:
            trace.add(
                "image_failed_directly",
                "_build_effective_question",
                uses={"attachments_count": len(body.attachments or [])},
                result={"error": image_analysis.get("error", "")},
                note="用户只发了图片但图片识别失败，因此不会进入知识库检索和模型回答。",
            )
            trace.finish("failed")
            db.close()

            async def failure_stream():
                for payload in _trace_sse_payloads(trace):
                    yield payload
                analysis_data = json.dumps(
                    {
                        "type": "image_analysis",
                        "analysis": image_analysis,
                    },
                    ensure_ascii=False,
                )
                yield f"data: {analysis_data}\n\n"
                error_data = json.dumps(
                    {
                        "type": "error",
                        "message": image_analysis.get("error") or "图片内容识别失败，请检查清晰度后重新上传。",
                    },
                    ensure_ascii=False,
                )
                yield f"data: {error_data}\n\n"
                yield "data: [DONE]\n\n"

            return StreamingResponse(failure_stream(), media_type="text/event-stream")
        conv = db.query(Conversation).filter_by(id=cid, user_id=user.id).first() if cid else None
        if not conv:
            # create new conversation
            cid = _new_id()
            title_source = raw_question or effective_question or display_question
            title = title_source[:30] + ("..." if len(title_source) > 30 else "")
            conv = Conversation(
                id=cid,
                user_id=user.id,
                knowledge_base_id=knowledge_base.id,
                title=title,
            )
            db.add(conv)
            db.commit()
            trace.add(
                "conversation_created",
                "stream_chat",
                creates={"conversation_id": cid, "title": title},
                result={"knowledge_base_id": knowledge_base.id},
                note="这是新对话，系统创建 conversation，并把它绑定到当前知识库。",
            )
        else:
            knowledge_base = conv.knowledge_base or knowledge_base
            trace.add(
                "conversation_loaded",
                "stream_chat",
                uses={"conversation_id": cid},
                result={"knowledge_base_id": knowledge_base.id, "title": conv.title},
                note="这是已有对话，系统复用它原本绑定的知识库，避免会话中途串库。",
            )
        trace.attach(conversation_id=cid)

        # save user message
        user_message = Message(
            conversation_id=cid,
            role="user",
            content=display_question,
            attachments=json.dumps(body.attachments, ensure_ascii=False),
        )
        db.add(user_message)
        db.commit()
        db.refresh(user_message)
        trace.add(
            "user_message_saved",
            "Message",
            creates={"user_message_id": user_message.id},
            params={"content": display_question, "attachments_count": len(body.attachments or [])},
            note="用户消息先写入数据库，后面的滑动窗口会排除这条当前消息，避免重复塞进 prompt。",
        )

        memory_context = _build_memory_context(
            db,
            conv,
            current_message_id=user_message.id,
        )
        retrieval_question = _build_memory_aware_retrieval_question(effective_question, memory_context)
        trace.add(
            "memory_built",
            "_build_memory_context",
            uses={
                "conversation_id": cid,
                "memory_summary": conv.memory_summary or "",
                "summary_upto_message_id": conv.memory_summary_upto_message_id or 0,
            },
            creates={
                "memory_context": memory_context,
                "retrieval_question": retrieval_question,
            },
            result={
                "memory_used": bool(memory_context),
                "used_for_retrieval": retrieval_question != effective_question,
                "window_turns": MEMORY_WINDOW_TURNS,
            },
            note="系统构造短期滑动窗口和长期摘要记忆，并生成真正用于 RAG 检索的问题。",
        )

        # retrieve relevant knowledge
        knowledge_chunks, retrieval_trace = await retrieve_knowledge(
            retrieval_question,
            knowledge_base_id=knowledge_base.id,
            db=db,
        )
        retrieval_trace = retrieval_trace or {}
        retrieval_trace["memory"] = {
            "used": bool(memory_context),
            "used_for_retrieval": retrieval_question != effective_question,
            "window_turns": MEMORY_WINDOW_TURNS,
            "summary_available": bool(conv.memory_summary),
            "summary_upto_message_id": conv.memory_summary_upto_message_id or 0,
        }
        if body.attachments:
            retrieval_trace["image_analysis_status"] = image_analysis.get("status", "")
            retrieval_trace["image_analysis_error"] = image_analysis.get("error", "")
            retrieval_trace["image_description"] = image_analysis.get("description", "")
            retrieval_trace["effective_question"] = effective_question
        trace.add(
            "retrieval_completed",
            "retrieve_knowledge",
            params={"question": retrieval_question, "knowledge_base_id": knowledge_base.id},
            creates={
                "query_plan": retrieval_trace.get("query_plan", {}),
                "routes": retrieval_trace.get("routes", []),
                "rrf": retrieval_trace.get("rrf", []),
                "rerank": retrieval_trace.get("rerank", {}),
            },
            result={"final_chunks_count": len(knowledge_chunks)},
            note="RAG 检索完成：包括问题规划、多路召回、RRF 融合、LLM rerank 和最终上下文选择。",
        )
        context = ""
        sources = _build_sources(knowledge_chunks)
        retrieved_contexts = [c.get("content", "") for c in knowledge_chunks if c.get("content")]
        if knowledge_chunks:
            context = "\n\n".join(
                f"[鏉ユ簮: {c['file_name']}]\n{c['content']}"
                for c in knowledge_chunks
            )
        trace.add(
            "context_built",
            "_build_sources",
            creates={"context": context, "sources": sources},
            result={"sources_count": len(sources), "retrieved_contexts_count": len(retrieved_contexts)},
            note="系统把最终选中的 chunk 拼成给大模型看的知识库上下文，并生成前端可展开的参考资料。",
        )

        async def event_stream():
            full = ""
            failed = False
            first_chunk_seen = False
            conversation_data = json.dumps({
                "type": "conversation",
                "conversation": {
                    "id": cid,
                    "title": conv.title,
                    "knowledge_base_id": knowledge_base.id,
                    "knowledge_base_name": knowledge_base.name,
                },
            }, ensure_ascii=False)
            if body.attachments:
                analysis_data = json.dumps(
                    {
                        "type": "image_analysis",
                        "analysis": image_analysis,
                    },
                    ensure_ascii=False,
                )
                for payload in _trace_sse_payloads(trace):
                    yield payload
                yield f"data: {analysis_data}\n\n"
            for payload in _trace_sse_payloads(trace):
                yield payload
            yield f"data: {conversation_data}\n\n"
            if sources:
                data = json.dumps({"type": "sources", "sources": sources}, ensure_ascii=False)
                yield f"data: {data}\n\n"
            trace.add(
                "generation_started",
                "_stream_deepseek_response",
                params={
                    "question": effective_question,
                    "memory_context": memory_context,
                    "context": context,
                },
                note="开始调用文本模型。会话记忆只辅助理解指代，事实依据仍优先来自知识库上下文。",
            )
            for payload in _trace_sse_payloads(trace):
                yield payload
            async for event in _stream_deepseek_response(effective_question, context, memory_context, trace):
                for payload in _trace_sse_payloads(trace):
                    yield payload
                if isinstance(event, dict):
                    if event.get("type") == "error":
                        failed = True
                        trace.add(
                            "generation_failed",
                            "_stream_deepseek_response",
                            result={"message": event.get("message") or event.get("content") or "DeepSeek 网络请求失败"},
                            note="模型生成阶段失败，系统会返回错误事件，并且不会保存失败 assistant 消息。",
                        )
                        for payload in _trace_sse_payloads(trace):
                            yield payload
                        data = json.dumps(
                            {
                                "type": "error",
                                "message": event.get("message") or event.get("content") or "DeepSeek 网络请求失败",
                            },
                            ensure_ascii=False,
                        )
                        yield f"data: {data}\n\n"
                        break
                    chunk = event.get("content", "")
                else:
                    chunk = event
                data = json.dumps({"content": chunk}, ensure_ascii=False)
                yield f"data: {data}\n\n"
                if not failed:
                    if chunk and not first_chunk_seen:
                        first_chunk_seen = True
                        trace.add(
                            "first_content_chunk",
                            "_stream_openai_chat_chunks",
                            result={"chunk": chunk},
                            note="大模型开始返回第一段流式内容，前端会逐步拼接为正在生成的回答。",
                        )
                        for payload in _trace_sse_payloads(trace):
                            yield payload
                    full += chunk

            # save assistant message
            if full and not failed:
                _safe_trace_add(
                    trace,
                    "assistant_ready_to_save",
                    "Message",
                    uses={"full_answer": full, "sources_count": len(sources)},
                    note="模型完整回答成功，系统准备保存 assistant 消息，并启动 RAGAS 和摘要判断。",
                )
                retrieval_trace["learning_trace"] = compact_trace_reference(trace.snapshot())
                assistant_message = None
                try:
                    assistant_message = Message(
                        conversation_id=cid,
                        role="assistant",
                        content=full,
                        sources=json.dumps(sources, ensure_ascii=False),
                        ragas_status="pending",
                        retrieval_trace=json.dumps(retrieval_trace, ensure_ascii=False),
                    )
                    db.add(assistant_message)
                    db.commit()
                    db.refresh(assistant_message)
                except Exception as exc:
                    db.rollback()
                    logger.warning("Assistant message save failed after stream finished: %s", exc, exc_info=True)
                    _safe_trace_add(
                        trace,
                        "assistant_save_failed",
                        "Message",
                        result={"error": str(exc)},
                        note="模型回答已经生成完毕，但保存 assistant 消息失败。系统仍会结束流，避免前端误报 network error。",
                    )

                if assistant_message:
                    _safe_trace_attach(trace, conversation_id=cid, message_id=assistant_message.id)
                    _safe_trace_add(
                        trace,
                        "assistant_message_saved",
                        "Message",
                        creates={"assistant_message_id": assistant_message.id},
                        result={"ragas_status": "pending"},
                        note="assistant 消息保存成功，历史会话刷新后仍可从这条消息打开流程。",
                    )
                    try:
                        schedule_ragas_evaluation(
                            assistant_message.id,
                            effective_question,
                            full,
                            retrieved_contexts,
                            trace.trace_id,
                        )
                        _safe_trace_add(
                            trace,
                            "ragas_scheduled",
                            "schedule_ragas_evaluation",
                            params={
                                "message_id": assistant_message.id,
                                "question": effective_question,
                                "answer_chars": len(full),
                                "contexts_count": len(retrieved_contexts),
                            },
                            note="RAGAS 在 assistant 保存后异步启动，不阻塞用户看到答案。",
                        )
                    except Exception as exc:
                        logger.warning("RAGAS schedule failed after stream finished: %s", exc, exc_info=True)
                        _safe_trace_add(
                            trace,
                            "ragas_schedule_failed",
                            "schedule_ragas_evaluation",
                            result={"error": str(exc)},
                            note="RAGAS 调度失败，但不影响主回答完成。",
                        )
                    try:
                        _schedule_memory_summary_update(cid, trace.trace_id)
                        _safe_trace_add(
                            trace,
                            "memory_summary_check_scheduled",
                            "_schedule_memory_summary_update",
                            params={"conversation_id": cid},
                            note="系统异步检查是否满足摘要压缩条件：完整回答轮数超过 8 且未摘要轮数至少 4。",
                        )
                    except Exception as exc:
                        logger.warning("Memory summary schedule failed after stream finished: %s", exc, exc_info=True)
                        _safe_trace_add(
                            trace,
                            "memory_summary_schedule_failed",
                            "_schedule_memory_summary_update",
                            result={"error": str(exc)},
                            note="摘要调度失败，但不影响主回答完成。",
                        )
                    try:
                        retrieval_trace["learning_trace"] = compact_trace_reference(trace.snapshot())
                        assistant_message.retrieval_trace = json.dumps(retrieval_trace, ensure_ascii=False)
                        db.commit()
                    except Exception as exc:
                        db.rollback()
                        logger.warning("Assistant trace reference update failed: %s", exc, exc_info=True)
                _safe_trace_finish(
                    trace,
                    "done" if assistant_message else "partial",
                    conversation_id=cid,
                    message_id=assistant_message.id if assistant_message else None,
                )
                for payload in _trace_sse_payloads(trace):
                    yield payload
            elif failed:
                _safe_trace_add(
                    trace,
                    "assistant_not_saved",
                    "Message",
                    result={"saved": False},
                    note="回答生成失败，遵循项目规则：不把失败内容保存为正式 assistant 消息。",
                )
                _safe_trace_finish(trace, "failed", conversation_id=cid)
                for payload in _trace_sse_payloads(trace):
                    yield payload
            db.close()
            yield "data: [DONE]\n\n"

        return StreamingResponse(event_stream(), media_type="text/event-stream")
    except Exception:
        db.close()
        raise



def _build_sources(chunks: list[dict]) -> list[dict]:
    sources = []
    for index, chunk in enumerate(chunks, start=1):
        content = (chunk.get("content") or "").strip()
        if not content:
            continue
        sources.append({
            "index": index,
            "file_id": chunk.get("file_id", 0),
            "file_name": chunk.get("file_name", "Untitled"),
            "chunk_id": chunk.get("chunk_id", ""),
            "route": chunk.get("route", ""),
            "routes": chunk.get("routes", []),
            "rrf_score": chunk.get("rrf_score"),
            "rerank_score": chunk.get("rerank_score"),
            "rerank_reason": chunk.get("rerank_reason", ""),
            "content": content,
            "excerpt": content[:180] + ("..." if len(content) > 180 else ""),
        })
    return sources


def _loads_json(value: str, default):
    if not value:
        return default
    try:
        return json.loads(value)
    except (TypeError, json.JSONDecodeError):
        return default


def _trace_sse_payloads(trace: TraceRecorder) -> list[str]:
    return [
        f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"
        for payload in trace.drain_sse_payloads()
    ]


def _safe_trace_add(trace: TraceRecorder, *args, **kwargs):
    try:
        return trace.add(*args, **kwargs)
    except Exception as exc:
        logger.warning("Learning trace add failed: %s", exc, exc_info=True)
        return {}


def _safe_trace_finish(trace: TraceRecorder, *args, **kwargs):
    try:
        trace.finish(*args, **kwargs)
    except Exception as exc:
        logger.warning("Learning trace finish failed: %s", exc, exc_info=True)


def _safe_trace_attach(trace: TraceRecorder, *args, **kwargs):
    try:
        trace.attach(*args, **kwargs)
    except Exception as exc:
        logger.warning("Learning trace attach failed: %s", exc, exc_info=True)


def _serialize_message(message: Message) -> dict:
    retrieval_trace = _loads_json(message.retrieval_trace, {})
    if not isinstance(retrieval_trace, dict):
        retrieval_trace = {}

    return {
        "id": message.id,
        "role": message.role,
        "content": message.content,
        "sources": _loads_json(message.sources, []),
        "attachments": _loads_json(message.attachments, []),
        "ragas_status": message.ragas_status or "",
        "ragas_scores": _loads_json(message.ragas_scores, {}),
        "ragas_error": message.ragas_error or "",
        "retrieval_trace": retrieval_trace,
        "image_analysis_status": retrieval_trace.get("image_analysis_status", ""),
        "image_analysis_error": retrieval_trace.get("image_analysis_error", ""),
        "image_description": retrieval_trace.get("image_description", ""),
        "created_at": message.created_at.isoformat() if message.created_at else "",
    }


def _build_memory_context(db: Session, conversation: Conversation, current_message_id: int) -> str:
    summary = (conversation.memory_summary or "").strip()
    summary_upto = conversation.memory_summary_upto_message_id or 0
    recent_messages = (
        db.query(Message)
        .filter(
            Message.conversation_id == conversation.id,
            Message.id < current_message_id,
            Message.id > summary_upto,
        )
        .order_by(Message.id.desc())
        .limit(max(MEMORY_WINDOW_TURNS, 1) * 2)
        .all()
    )
    recent_text = _format_recent_memory_messages(list(reversed(recent_messages)))

    sections = []
    if summary:
        sections.append(f"长期摘要记忆：\n{_clip_text(summary, MEMORY_SUMMARY_MAX_CHARS)}")
    if recent_text:
        sections.append(f"最近对话窗口：\n{recent_text}")
    return "\n\n".join(sections).strip()


def _build_memory_aware_retrieval_question(question: str, memory_context: str) -> str:
    if not memory_context:
        return question
    compact_memory = _clip_text(memory_context, 1500)
    return (
        "以下会话记忆仅用于消解当前问题中的指代和省略，不作为事实依据。\n"
        f"{compact_memory}\n\n"
        f"当前检索问题：{question}"
    )


def _format_recent_memory_messages(messages: list[Message]) -> str:
    lines = []
    for message in messages:
        content = (message.content or "").strip()
        if not content:
            continue
        role = "用户" if message.role == "user" else "助手"
        lines.append(f"{role}：{content}")

    selected = []
    total = 0
    for line in reversed(lines):
        line_length = len(line)
        if selected and total + line_length > MEMORY_RECENT_MAX_CHARS:
            break
        if line_length > MEMORY_RECENT_MAX_CHARS:
            line = _clip_text(line, MEMORY_RECENT_MAX_CHARS)
            line_length = len(line)
        selected.append(line)
        total += line_length
    return "\n".join(reversed(selected))


def _schedule_memory_summary_update(conversation_id: str, trace_id: str | None = None):
    try:
        asyncio.create_task(_maybe_update_memory_summary(conversation_id, trace_id))
    except RuntimeError:
        logger.warning("Unable to schedule memory summary update: conversation_id=%s", conversation_id)
        append_trace_event(
            trace_id,
            "memory_summary_schedule_failed",
            "_schedule_memory_summary_update",
            result={"conversation_id": conversation_id},
            note="当前事件循环不可用，摘要压缩检查没有成功启动。",
        )


async def _maybe_update_memory_summary(conversation_id: str, trace_id: str | None = None):
    db = SessionLocal()
    try:
        conversation = db.query(Conversation).filter_by(id=conversation_id).first()
        if not conversation:
            append_trace_event(
                trace_id,
                "memory_summary_skipped",
                "_maybe_update_memory_summary",
                result={"reason": "conversation_not_found"},
                note="摘要检查跳过：会话不存在。",
            )
            return

        assistant_messages = (
            db.query(Message)
            .filter_by(conversation_id=conversation_id, role="assistant")
            .order_by(Message.id.asc())
            .all()
        )
        if len(assistant_messages) <= MEMORY_SUMMARY_TRIGGER_TURNS:
            append_trace_event(
                trace_id,
                "memory_summary_skipped",
                "_maybe_update_memory_summary",
                result={
                    "assistant_turns": len(assistant_messages),
                    "trigger_turns": MEMORY_SUMMARY_TRIGGER_TURNS,
                    "reason": "turns_not_enough",
                },
                note="摘要检查跳过：完整回答轮数还没有超过触发阈值。",
            )
            return

        cutoff_index = len(assistant_messages) - max(MEMORY_WINDOW_TURNS, 1) - 1
        if cutoff_index < 0:
            append_trace_event(
                trace_id,
                "memory_summary_skipped",
                "_maybe_update_memory_summary",
                result={"reason": "no_window_outside_messages"},
                note="摘要检查跳过：窗口外没有足够历史消息可压缩。",
            )
            return
        cutoff_message = assistant_messages[cutoff_index]
        summary_upto = conversation.memory_summary_upto_message_id or 0
        unsummarized_turns = [
            message for message in assistant_messages
            if summary_upto < message.id <= cutoff_message.id
        ]
        if len(unsummarized_turns) < MEMORY_WINDOW_TURNS:
            append_trace_event(
                trace_id,
                "memory_summary_skipped",
                "_maybe_update_memory_summary",
                result={
                    "unsummarized_turns": len(unsummarized_turns),
                    "required_turns": MEMORY_WINDOW_TURNS,
                    "summary_upto_message_id": summary_upto,
                },
                note="摘要检查跳过：距离上次摘要之后新增的窗口外轮数不足，避免每轮都调用模型摘要。",
            )
            return

        messages_to_summarize = (
            db.query(Message)
            .filter(
                Message.conversation_id == conversation_id,
                Message.id > summary_upto,
                Message.id <= cutoff_message.id,
            )
            .order_by(Message.id.asc())
            .all()
        )
        transcript = _format_messages_for_summary(messages_to_summarize)
        if not transcript:
            append_trace_event(
                trace_id,
                "memory_summary_skipped",
                "_maybe_update_memory_summary",
                result={"reason": "empty_transcript"},
                note="摘要检查跳过：待摘要消息内容为空。",
            )
            return

        append_trace_event(
            trace_id,
            "memory_summary_triggered",
            "_maybe_update_memory_summary",
            params={
                "assistant_turns": len(assistant_messages),
                "trigger_turns": MEMORY_SUMMARY_TRIGGER_TURNS,
                "unsummarized_turns": len(unsummarized_turns),
                "cutoff_message_id": cutoff_message.id,
            },
            creates={"transcript": transcript},
            note="摘要条件满足，系统将窗口外历史和旧摘要合并压缩为新的长期记忆。",
        )
        next_summary = await _summarize_conversation_memory(
            conversation.memory_summary or "",
            transcript,
        )
        if not next_summary:
            append_trace_event(
                trace_id,
                "memory_summary_failed",
                "_summarize_conversation_memory",
                result={"reason": "empty_summary_or_model_failed"},
                note="摘要模型没有返回可用摘要，主回答不受影响，保留原摘要。",
            )
            return

        conversation.memory_summary = _clip_text(next_summary, MEMORY_SUMMARY_MAX_CHARS)
        conversation.memory_summary_upto_message_id = cutoff_message.id
        conversation.memory_updated_at = datetime.now()
        db.commit()
        append_trace_event(
            trace_id,
            "memory_summary_updated",
            "_summarize_conversation_memory",
            creates={
                "memory_summary": conversation.memory_summary,
                "memory_summary_upto_message_id": cutoff_message.id,
            },
            note="长期摘要已更新。后续追问会同时使用这份摘要和最近 4 轮滑动窗口。",
        )
        logger.info(
            "Conversation memory summary updated: conversation_id=%s upto_message_id=%s",
            conversation_id,
            cutoff_message.id,
        )
    except Exception as exc:
        logger.warning(
            "Conversation memory summary update failed: conversation_id=%s error=%s",
            conversation_id,
            exc,
            exc_info=True,
        )
        append_trace_event(
            trace_id,
            "memory_summary_failed",
            "_maybe_update_memory_summary",
            result={"error": str(exc)},
            note="摘要压缩发生异常，但它是异步辅助能力，不会影响本轮回答。",
        )
    finally:
        db.close()


def _format_messages_for_summary(messages: list[Message]) -> str:
    lines = []
    for message in messages:
        content = _clip_text((message.content or "").strip(), 1200)
        if not content:
            continue
        role = "用户" if message.role == "user" else "助手"
        lines.append(f"{role}：{content}")
    return "\n".join(lines)


async def _summarize_conversation_memory(previous_summary: str, transcript: str) -> str:
    if not DEEPSEEK_API_KEY:
        logger.warning("Skip memory summary update because DeepSeek API Key is not configured")
        return ""

    system_prompt = (
        "你是企业知识库问答系统的会话记忆压缩器。"
        "请把旧摘要和新增对话压缩成可用于后续追问理解的中文摘要。"
        "只保留用户目标、当前任务背景、用户明确确认过的事实、关键约束/偏好/口径、未解决问题或待办。"
        "不要保存知识库原文大段内容，不要编造未确认事实，不要输出 Markdown 标题。"
    )
    user_prompt = (
        f"摘要长度上限：{MEMORY_SUMMARY_MAX_CHARS} 字。\n\n"
        f"旧摘要：\n{previous_summary or '（无）'}\n\n"
        f"新增对话：\n{transcript}\n\n"
        "请输出新的会话摘要："
    )
    payload = {
        "model": normalize_deepseek_model(DEEPSEEK_MODEL),
        "stream": False,
        "temperature": 0.1,
        "max_tokens": 900,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    }
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json",
    }
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            response = await client.post(deepseek_chat_url(), json=payload, headers=headers)
        if response.status_code >= 400:
            logger.warning(
                "Memory summary call failed: status=%s detail=%s",
                response.status_code,
                response.text[:300],
            )
            return ""
        content = response.json().get("choices", [{}])[0].get("message", {}).get("content", "")
        return _clip_text(content.strip(), MEMORY_SUMMARY_MAX_CHARS)
    except Exception as exc:
        logger.warning("Memory summary network request failed: %s", exc, exc_info=True)
        return ""


def _clip_text(text: str, max_chars: int) -> str:
    value = (text or "").strip()
    if max_chars <= 0 or len(value) <= max_chars:
        return value
    return value[:max_chars].rstrip() + "..."


def _ensure_oss_config():
    if not all([OSS_ACCESS_KEY_ID, OSS_ACCESS_KEY_SECRET, OSS_BUCKET, OSS_ENDPOINT]):
        raise HTTPException(500, "OSS 环境变量未完整配置")


def _oss_host() -> str:
    endpoint = OSS_ENDPOINT.replace("https://", "").replace("http://", "").rstrip("/")
    if endpoint.startswith(f"{OSS_BUCKET}."):
        return endpoint
    return f"{OSS_BUCKET}.{endpoint}"


def _oss_object_path(object_key: str) -> str:
    return "/" + quote(object_key, safe="/")


def _oss_signature(string_to_sign: str) -> str:
    digest = hmac.new(
        OSS_ACCESS_KEY_SECRET.encode("utf-8"),
        string_to_sign.encode("utf-8"),
        hashlib.sha1,
    ).digest()
    return base64.b64encode(digest).decode("utf-8")


async def _put_oss_object(object_key: str, content: bytes, content_type: str):
    _ensure_oss_config()
    host = _oss_host()
    date = formatdate(usegmt=True)
    resource = f"/{OSS_BUCKET}/{object_key}"
    string_to_sign = f"PUT\n\n{content_type}\n{date}\nx-oss-object-acl:public-read\n{resource}"
    signature = _oss_signature(string_to_sign)
    url = f"https://{host}{_oss_object_path(object_key)}"
    headers = {
        "Authorization": f"OSS {OSS_ACCESS_KEY_ID}:{signature}",
        "Content-Type": content_type,
        "Date": date,
        "Host": host,
        "x-oss-object-acl": "public-read",
    }

    async with httpx.AsyncClient(timeout=60) as client:
        response = await client.put(url, content=content, headers=headers)
    if response.status_code >= 400:
        raise RuntimeError(f"{response.status_code} {response.text[:200]}")


def _sign_oss_url(object_key: str, expires: int = 3600) -> str:
    _ensure_oss_config()
    expires_at = int(time.time()) + expires
    resource = f"/{OSS_BUCKET}/{object_key}"
    string_to_sign = f"GET\n\n\n{expires_at}\n{resource}"
    signature = _oss_signature(string_to_sign)
    query = urlencode({
        "OSSAccessKeyId": OSS_ACCESS_KEY_ID,
        "Expires": str(expires_at),
        "Signature": signature,
    })
    return f"https://{_oss_host()}{_oss_object_path(object_key)}?{query}"


def _public_oss_url(object_key: str) -> str:
    _ensure_oss_config()
    return f"https://{_oss_host()}{_oss_object_path(object_key)}"



def _openai_chat_url(base_url: str) -> str:
    base_url = (base_url or "").rstrip("/")
    if base_url.endswith("/v1"):
        return f"{base_url}/chat/completions"
    return f"{base_url}/v1/chat/completions"


async def _build_effective_question(
    question: str,
    attachments: Optional[list[dict]] = None,
) -> tuple[str, dict]:
    attachments = attachments or []
    question = (question or "").strip()
    if not attachments:
        return question, {}

    image_analysis = await _analyze_image_attachments(attachments, question)
    description = image_analysis.get("description", "").strip()

    if question and description:
        return f"{question}\n\n图片内容：{description}", image_analysis
    if description:
        return description, image_analysis
    return question, image_analysis


async def _analyze_image_attachments(attachments: list[dict], question: str = "") -> dict:
    image_urls = _build_image_urls(attachments)
    if not image_urls:
        return {
            "status": "failed",
            "description": "",
            "error": "图片附件缺少可访问的 OSS object_key，请重新上传后再试。",
        }
    if not VISION_API_KEY:
        return {
            "status": "failed",
            "description": "",
            "error": "图片内容提取失败，请检查 VISION_MODEL/VISION_API_KEY/OSS URL。",
        }

    prompts = _image_analysis_prompts(question)
    last_error = ""
    last_description = ""
    for prompt in prompts:
        try:
            description = await _request_image_description(prompt, image_urls)
        except Exception as exc:
            last_error = str(exc)
            continue

        description = (description or "").strip()
        if not description:
            last_error = "图片内容提取失败：视觉模型未返回图片描述。"
            continue

        status, error = _classify_image_analysis(description)
        if status == "failed":
            last_error = error or "图片内容识别失败，请检查清晰度后重新上传。"
            last_description = description
            continue

        return {
            "status": status,
            "description": description,
            "error": error,
        }

    if last_description:
        status, error = _classify_image_analysis(last_description)
        if status in {"partial", "success"}:
            return {
                "status": status,
                "description": last_description,
                "error": error,
            }

    return {
        "status": "failed",
        "description": "",
        "error": last_error or "图片内容识别失败，请检查清晰度后重新上传。",
    }


def _image_analysis_prompts(question: str = "") -> list[str]:
    question = (question or "").strip()
    if question:
        return [
            (
                "请先逐字识别图片中可见的文字、题目、选项、表格、数字和图表标签，"
                "再用中文概括图片内容，并明确说明哪些内容看清了，哪些内容不够清晰。"
                "不要编造图片中不存在的信息。"
                f"\n用户问题：{question}\n请重点关注与该问题相关的图片信息。"
            ),
            (
                "请再次仔细查看这张图片，优先识别可见文字、题目、选项、数字和表格。"
                "如果局部模糊，也要尽量保留可辨认部分，并说明无法确认的区域。"
                "不要编造图片中不存在的信息。"
                f"\n用户问题：{question}"
            ),
        ]
    return [
        (
            "请尽可能逐字识别这张图片中的所有可见文字、题目、选项、数字、表格和图表标签，"
            "先抄录能看清的内容，再总结图片整体含义。"
            "如果有模糊或看不清的部分，请明确指出，不要编造。"
        ),
        (
            "请重新检查这张图片，专注于文字识别和局部细节。"
            "请输出可辨认的原文、题目、选项、数字和表格内容，并说明哪些地方无法识别。"
        ),
    ]


async def _request_image_description(prompt: str, image_urls: list[str]) -> str:
    user_content = [{"type": "text", "text": prompt}]
    user_content.extend({"type": "image_url", "image_url": {"url": url}} for url in image_urls)
    payload = {
        "model": VISION_MODEL,
        "stream": False,
        "messages": [
            {"role": "system", "content": "你是图片内容理解助手，负责将图片转写为适合检索和问答的中文文本。"},
            {"role": "user", "content": user_content},
        ],
    }
    headers = {
        "Authorization": f"Bearer {VISION_API_KEY}",
        "Content-Type": "application/json",
    }
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            response = await client.post(_openai_chat_url(VISION_BASE_URL), json=payload, headers=headers)
        if response.status_code >= 400:
            detail = response.text[:300]
            logger.warning("Vision description failed: status=%s detail=%s", response.status_code, detail)
            raise RuntimeError("图片内容提取失败，请检查 VISION_MODEL/VISION_API_KEY/OSS URL。")
        description = response.json().get("choices", [{}])[0].get("message", {}).get("content", "")
        return (description or "").strip()
    except httpx.HTTPError as exc:
        logger.warning("Vision description network request failed: %s", exc, exc_info=True)
        raise RuntimeError("图片内容提取失败，请检查 VISION_MODEL/VISION_API_KEY/OSS URL。") from exc


def _classify_image_analysis(description: str) -> tuple[str, str]:
    text = (description or "").strip()
    if not text:
        return "failed", "图片内容提取失败：视觉模型未返回图片描述。"

    failure_phrases = [
        "看不清",
        "无法识别",
        "无法看清",
        "无法辨认",
        "无法读取",
        "过于模糊",
        "图片模糊",
        "不清晰",
        "无法确定",
        "难以识别",
        "图片质量",
        "无法完全识别",
    ]
    partial_phrases = [
        "部分",
        "大致",
        "可能",
        "疑似",
        "不够清晰",
        "部分可见",
        "仅能看到",
        "只能识别",
    ]
    useful_markers = [
        "可见",
        "看到",
        "能看到",
        "识别到",
        "文字为",
        "内容为",
        "显示",
        "题目",
        "选项",
        "数字",
        "表格",
        "图表",
    ]

    if any(phrase in text for phrase in failure_phrases):
        has_partial_signal = any(phrase in text for phrase in partial_phrases)
        has_useful_content = any(marker in text for marker in useful_markers)
        if has_partial_signal and has_useful_content:
            return "partial", "图片内容仅部分识别，请结合文字问题查看。"
        return "failed", "图片内容未能清晰识别，请检查图片清晰度后重试。"

    if any(phrase in text for phrase in partial_phrases):
        return "partial", "图片内容仅部分识别，请结合文字问题查看。"

    return "success", ""


async def _stream_deepseek_response(
    question: str,
    context: str,
    memory_context: str = "",
    trace: TraceRecorder | None = None,
):
    if not DEEPSEEK_API_KEY:
        if trace:
            trace.add(
                "deepseek_missing_config",
                "_stream_deepseek_response",
                result={"configured": False},
                note="DeepSeek API Key 未配置，文本生成无法开始。",
            )
        yield {"type": "error", "message": _model_missing_error(False)}
        return

    system_prompt = (
        "你是企业知识库智能问答助手。请优先基于提供的企业知识库上下文回答用户问题；"
        "如果知识库信息不足，请明确说明不足之处，并给出可验证的建议。"
        "回答使用中文，结构清晰，避免编造未在上下文中出现的事实。"
        "信息优先级为：企业知识库上下文 > 当前用户问题 > 会话记忆。"
        "会话记忆只能用于理解指代、延续任务和用户偏好，不能替代知识库事实依据。"
    )
    user_prompt = (
        f"用户问题：{question}\n\n"
        f"会话记忆：\n{memory_context or '（无）'}\n\n"
        f"企业知识库上下文：\n{context or '（未检索到相关知识库内容）'}"
    )
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    if trace:
        trace.add(
            "generation_prompt_built",
            "_stream_deepseek_response",
            creates={
                "system_prompt": system_prompt,
                "user_prompt": user_prompt,
            },
            note="系统构造最终给文本模型的 messages。注意：会话记忆和知识库上下文是分区放入 prompt 的。",
        )

    async for event in _stream_openai_chat_chunks(
        deepseek_chat_url(),
        DEEPSEEK_API_KEY,
        normalize_deepseek_model(DEEPSEEK_MODEL),
        messages,
        "DeepSeek",
    ):
        if isinstance(event, dict) and event.get("type") == "error":
            if trace:
                trace.add(
                    "deepseek_failed",
                    "_stream_openai_chat_chunks",
                    result={
                        "reason": event.get("reason"),
                        "status_code": event.get("status_code"),
                        "detail": event.get("detail", ""),
                    },
                    note="DeepSeek 流式调用失败。若属于网络或 5xx 错误，系统会尝试文本后备模型。",
                )
            if _should_use_text_fallback(event):
                logger.warning(
                    "DeepSeek failed with retryable error, switching to text fallback: reason=%s status=%s",
                    event.get("reason"),
                    event.get("status_code"),
                )
                if trace:
                    trace.add(
                        "text_fallback_started",
                        "_stream_text_fallback_response",
                        params={"model": TEXT_FALLBACK_MODEL, "enabled": TEXT_FALLBACK_ENABLED},
                        note="DeepSeek 暂不可用，系统尝试使用兼容文本后备模型继续回答。",
                    )
                async for fallback_event in _stream_text_fallback_response(messages, trace):
                    yield fallback_event
                return
            yield {
                "type": "error",
                "message": _model_error_message(
                    False,
                    event.get("reason") or "api",
                    event.get("status_code"),
                    DEEPSEEK_MODEL,
                ),
            }
            return
        yield event


async def _stream_text_fallback_response(messages: list[dict], trace: TraceRecorder | None = None):
    if not TEXT_FALLBACK_ENABLED:
        if trace:
            trace.add(
                "text_fallback_disabled",
                "_stream_text_fallback_response",
                result={"enabled": False},
                note="文本后备模型未启用，因此 DeepSeek 失败后无法继续生成。",
            )
        yield {
            "type": "error",
            "message": "DeepSeek 网络请求失败，且文本后备模型未启用，请设置 TEXT_FALLBACK_ENABLED=true。",
        }
        return
    if not TEXT_FALLBACK_API_KEY:
        if trace:
            trace.add(
                "text_fallback_missing_config",
                "_stream_text_fallback_response",
                result={"configured": False},
                note="文本后备模型缺少 API Key，无法接管 DeepSeek 的失败请求。",
            )
        yield {
            "type": "error",
            "message": "DeepSeek 网络请求失败，且文本后备模型未配置，请配置 TEXT_FALLBACK_API_KEY 或 DASHSCOPE_API_KEY。",
        }
        return

    async for event in _stream_openai_chat_chunks(
        _openai_chat_url(TEXT_FALLBACK_BASE_URL),
        TEXT_FALLBACK_API_KEY,
        TEXT_FALLBACK_MODEL,
        messages,
        "百炼文本后备模型",
    ):
        if isinstance(event, dict) and event.get("type") == "error":
            if trace:
                trace.add(
                    "text_fallback_failed",
                    "_stream_openai_chat_chunks",
                    result={
                        "reason": event.get("reason"),
                        "status_code": event.get("status_code"),
                        "detail": event.get("detail", ""),
                    },
                    note="文本后备模型也调用失败，本轮回答会以错误结束。",
                )
            yield {
                "type": "error",
                "message": _text_fallback_error_message(
                    event.get("reason") or "api",
                    event.get("status_code"),
                ),
            }
            return
        if trace:
            trace.add(
                "text_fallback_chunk",
                "_stream_openai_chat_chunks",
                result={"chunk": event},
                note="文本后备模型开始返回内容，前端会继续按同一条流式回答展示。",
            )
            trace = None
        yield event


async def _stream_openai_chat_chunks(
    url: str,
    api_key: str,
    model: str,
    messages: list[dict],
    provider_label: str,
):
    payload = {
        "model": model,
        "stream": True,
        "messages": messages,
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    try:
        async with httpx.AsyncClient(timeout=60) as client:
            async with client.stream("POST", url, json=payload, headers=headers) as response:
                if response.status_code >= 400:
                    text = await response.aread()
                    detail = text.decode("utf-8", errors="replace")[:300]
                    logger.warning(
                        "%s call failed: status=%s detail=%s",
                        provider_label,
                        response.status_code,
                        detail,
                    )
                    yield {
                        "type": "error",
                        "reason": "api",
                        "status_code": response.status_code,
                        "detail": detail,
                    }
                    return

                async for line in response.aiter_lines():
                    if not line.startswith("data:"):
                        continue
                    data = line.removeprefix("data:").strip()
                    if not data or data == "[DONE]":
                        continue
                    try:
                        parsed = json.loads(data)
                    except json.JSONDecodeError:
                        continue
                    delta = parsed.get("choices", [{}])[0].get("delta", {})
                    content = delta.get("content")
                    if content:
                        yield content
    except httpx.HTTPError as exc:
        logger.warning("%s network request failed: %s", provider_label, exc, exc_info=True)
        yield {
            "type": "error",
            "reason": "network",
            "status_code": None,
            "detail": str(exc),
        }


def _should_use_text_fallback(event: dict) -> bool:
    if event.get("reason") == "network":
        return True
    status_code = event.get("status_code")
    return isinstance(status_code, int) and 500 <= status_code < 600


def _text_fallback_error_message(reason: str, status_code: Optional[int] = None) -> str:
    if reason == "api":
        status = f"（HTTP {status_code}）" if status_code else ""
        return (
            f"DeepSeek 不可用，且百炼文本后备模型 {TEXT_FALLBACK_MODEL} 返回错误{status}。"
            "请检查 TEXT_FALLBACK_MODEL、TEXT_FALLBACK_API_KEY、TEXT_FALLBACK_BASE_URL。"
        )
    return (
        f"DeepSeek 不可用，且百炼文本后备模型 {TEXT_FALLBACK_MODEL} 网络请求失败。"
        "请检查 TEXT_FALLBACK_BASE_URL、网络连接和 DASHSCOPE_API_KEY。"
    )


def _model_missing_error(has_images: bool) -> str:
    if has_images:
        return "图片问答模型未配置，请配置 VISION_MODEL / VISION_API_KEY 后重试。"
    return "文本问答模型未配置，请配置 DEEPSEEK_MODEL / DEEPSEEK_API_KEY 后重试。"


def _model_error_message(
    has_images: bool,
    reason: str,
    status_code: Optional[int] = None,
    model: str = "",
) -> str:
    model_label = model or ("图片模型" if has_images else "文本模型")
    if reason == "api":
        status = f"（HTTP {status_code}）" if status_code else ""
        if has_images:
            return (
                f"图片问答生成失败：{model_label} 返回错误{status}。"
                "请检查 VISION_MODEL、VISION_API_KEY、VISION_BASE_URL 和 OSS 签名 URL 是否可访问。"
            )
        return (
            f"回答生成失败：{model_label} 返回错误{status}。"
            "请检查 DEEPSEEK_MODEL、DEEPSEEK_API_KEY、DEEPSEEK_BASE_URL 和网络连接。"
        )
    if has_images:
        return (
            f"图片问答生成失败：{model_label} 网络请求失败。"
            "请检查 VISION_MODEL、VISION_API_KEY、VISION_BASE_URL 和 OSS 签名 URL 是否可访问。"
        )
    return (
        f"回答生成失败：{model_label} 网络请求失败。"
        "请检查 DEEPSEEK_MODEL、DEEPSEEK_API_KEY、DEEPSEEK_BASE_URL 和网络连接。"
    )


def _build_image_urls(attachments: list[dict]) -> list[str]:
    urls = []
    for item in attachments:
        object_key = item.get("object_key")
        if object_key:
            urls.append(_public_oss_url(object_key))
    return urls


# ---------------------------------------------------------------------------
# knowledge endpoints
# ---------------------------------------------------------------------------


@app.get("/api/knowledge")
def list_knowledge(knowledge_base_id: Optional[int] = None,
                   _user: User = Depends(get_current_user),
                   db: Session = Depends(get_db)):
    knowledge_base = _resolve_knowledge_base(db, knowledge_base_id)
    files = (
        db.query(KnowledgeFile)
        .filter_by(knowledge_base_id=knowledge_base.id)
        .order_by(KnowledgeFile.created_at.desc())
        .all()
    )
    return [
        {
            "id": f.id,
            "knowledge_base_id": f.knowledge_base_id,
            "name": f.name,
            "size": f.size,
            "created_at": f.created_at.isoformat() if f.created_at else "",
        }
        for f in files
    ]


@app.post("/api/knowledge/upload")
async def upload_knowledge(file: UploadFile = File(...),
                           knowledge_base_id: Optional[int] = Form(None),
                           _user: User = Depends(get_current_user),
                           db: Session = Depends(get_db)):
    knowledge_base = _resolve_knowledge_base(db, knowledge_base_id)
    content = await file.read()
    text = _extract_file_text(file.filename or "", content)

    entry = KnowledgeFile(
        knowledge_base_id=knowledge_base.id,
        name=file.filename or "unknown",
        size=len(content),
        content=text,
    )
    db.add(entry)
    try:
        db.commit()
    except SQLAlchemyError as exc:
        db.rollback()
        logger.warning("Knowledge file metadata save failed: filename=%s error=%s", file.filename, exc, exc_info=True)
        raise HTTPException(500, _knowledge_file_save_error_message(exc))
    db.refresh(entry)

    chunks = _chunk_text(text, entry.id)
    from chroma_client import add_chunks, delete_file_chunks
    try:
        add_chunks(chunks, entry.id, entry.name, knowledge_base.id)
    except Exception as exc:
        logger.warning("Knowledge file indexing failed: file_id=%s error=%s", entry.id, exc, exc_info=True)
        try:
            delete_file_chunks(entry.id)
        except Exception as cleanup_exc:
            logger.warning(
                "Failed to clean partially indexed chunks: file_id=%s error=%s",
                entry.id,
                cleanup_exc,
                exc_info=True,
            )
        db.delete(entry)
        try:
            db.commit()
        except SQLAlchemyError as cleanup_commit_exc:
            db.rollback()
            logger.warning(
                "Failed to rollback partially indexed knowledge file: file_id=%s error=%s",
                entry.id,
                cleanup_commit_exc,
                exc_info=True,
            )
        raise HTTPException(500, "文件索引失败，已取消本次上传，请稍后重试")

    return {
        "id": entry.id,
        "knowledge_base_id": entry.knowledge_base_id,
        "name": entry.name,
        "size": entry.size,
        "created_at": entry.created_at.isoformat() if entry.created_at else "",
    }


@app.delete("/api/knowledge/{fid}")
def delete_knowledge(fid: int, _user: User = Depends(get_current_user),
                     db: Session = Depends(get_db)):
    entry = db.query(KnowledgeFile).filter_by(id=fid).first()
    if not entry:
        raise HTTPException(404, "文件不存在")
    from chroma_client import delete_file_chunks
    try:
        delete_file_chunks(fid)
    except Exception as exc:
        logger.warning("Knowledge file vector cleanup failed: file_id=%s error=%s", fid, exc, exc_info=True)
        raise HTTPException(500, "文件向量索引删除失败，已保留原文件，请稍后重试")
    db.delete(entry)
    db.commit()
    return {"message": "ok"}


@app.get("/api/knowledge/{fid}")
def get_knowledge_detail(fid: int, _user: User = Depends(get_current_user),
                         db: Session = Depends(get_db)):
    entry = db.query(KnowledgeFile).filter_by(id=fid).first()
    if not entry:
        raise HTTPException(404, "文件不存在")
    return {
        "id": entry.id,
        "knowledge_base_id": entry.knowledge_base_id,
        "name": entry.name,
        "size": entry.size,
        "created_at": entry.created_at.isoformat() if entry.created_at else "",
    }


@app.get("/api/knowledge/{fid}/content")
def get_knowledge_content(fid: int, _user: User = Depends(get_current_user),
                          db: Session = Depends(get_db)):
    entry = db.query(KnowledgeFile).filter_by(id=fid).first()
    if not entry:
        raise HTTPException(404, "文件不存在")
    content = entry.content or ""
    if not content and (entry.name or "").lower().endswith(".docx"):
        content = "该文件上传时未抽取内容，请重新上传以生成预览。"
    return {"id": entry.id, "name": entry.name, "content": content}


def _extract_file_text(filename: str, content: bytes) -> str:
    lower_name = filename.lower()
    if lower_name.endswith(".docx"):
        return _extract_docx_text(content)
    if lower_name.endswith(".pdf"):
        return _extract_pdf_text(content)
    return content.decode("utf-8", errors="replace")


def _extract_docx_text(content: bytes) -> str:
    from docx import Document

    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise HTTPException(400, f"DOCX 解析失败：{exc}")
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_pdf_text(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise HTTPException(500, "后端缺少 pypdf 依赖，无法解析 PDF")

    try:
        reader = PdfReader(BytesIO(content), strict=False)
        parts = []
        for index, page in enumerate(reader.pages, start=1):
            page_text = (page.extract_text() or "").strip()
            if page_text:
                parts.append(f"第 {index} 页\n{page_text}")
        return "\n\n".join(parts)
    except Exception as exc:
        raise HTTPException(400, f"PDF 解析失败：{exc}")


def _knowledge_file_save_error_message(exc: SQLAlchemyError) -> str:
    detail = str(exc)
    if "Incorrect string value" in detail or "1366" in detail:
        return (
            "文件内容包含中文字符，但当前 MySQL 表或字段仍不是 utf8mb4。"
            "请重启后端让启动迁移生效；如仍失败，请执行："
            "ALTER DATABASE rag_system CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci; "
            "ALTER TABLE knowledge_files CONVERT TO CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci; "
            "ALTER TABLE knowledge_files MODIFY COLUMN content LONGTEXT CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;"
        )
    return "文件信息写入数据库失败，请稍后重试"


def _chunk_text(text: str, file_id: int, chunk_size: int = 500, chunk_overlap: int = 50) -> list[dict]:
    """Split text into chunks with overlap and return them with ids."""
    if not text.strip():
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk_text = text[start:end].strip()
        if chunk_text:
            chunks.append({"id": f"{start}", "text": chunk_text})
        start += (chunk_size - chunk_overlap)
    return chunks


# ---------------------------------------------------------------------------
# checkpointer endpoints (for frontend state management)
# ---------------------------------------------------------------------------


@app.get("/api/checkpointer/threads")
def list_checkpointer_threads(_user: User = Depends(get_current_user)):
    return {"threads": list_threads()}


# ---------------------------------------------------------------------------
# entry
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8002, reload=True)

