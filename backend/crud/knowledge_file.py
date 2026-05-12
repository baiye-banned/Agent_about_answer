from io import BytesIO

from fastapi import HTTPException
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.orm import Session

from model.models import KnowledgeFile


def serialize_knowledge_file(file_entry: KnowledgeFile) -> dict:
    return {
        "id": file_entry.id,
        "knowledge_base_id": file_entry.knowledge_base_id,
        "name": file_entry.name,
        "size": file_entry.size,
        "created_at": file_entry.created_at.isoformat() if file_entry.created_at else "",
    }


def list_knowledge_files(db: Session, knowledge_base_id: int) -> list[KnowledgeFile]:
    return (
        db.query(KnowledgeFile)
        .filter_by(knowledge_base_id=knowledge_base_id)
        .order_by(KnowledgeFile.created_at.desc())
        .all()
    )


def get_knowledge_file(db: Session, fid: int) -> KnowledgeFile | None:
    return db.query(KnowledgeFile).filter_by(id=fid).first()


def create_knowledge_file(
    db: Session,
    *,
    knowledge_base_id: int,
    name: str,
    size: int,
    content: str,
) -> KnowledgeFile:
    entry = KnowledgeFile(
        knowledge_base_id=knowledge_base_id,
        name=name,
        size=size,
        content=content,
    )
    db.add(entry)
    db.commit()
    db.refresh(entry)
    return entry


def delete_knowledge_file(db: Session, fid: int) -> KnowledgeFile | None:
    entry = get_knowledge_file(db, fid)
    if not entry:
        return None
    db.delete(entry)
    db.commit()
    return entry


def get_knowledge_content(db: Session, fid: int) -> dict | None:
    entry = get_knowledge_file(db, fid)
    if not entry:
        return None
    content = entry.content or ""
    if not content and (entry.name or "").lower().endswith(".docx"):
        content = "该文件上传时未抽取内容，请重新上传以生成预览。"
    return {
        "id": entry.id,
        "name": entry.name,
        "content": content,
    }


def extract_file_text(filename: str, content: bytes) -> str:
    lower_name = filename.lower()
    if lower_name.endswith(".docx"):
        return extract_docx_text(content)
    if lower_name.endswith(".pdf"):
        return extract_pdf_text(content)
    return content.decode("utf-8", errors="replace")


def extract_docx_text(content: bytes) -> str:
    from docx import Document

    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise HTTPException(400, f"DOCX 解析失败：{exc}")
    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_pdf_text(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise HTTPException(500, "后端缺少 pypdf 依赖，无法解析 PDF")

    try:
        reader = PdfReader(BytesIO(content), strict=False)
        parts = []
        for index, page in enumerate(reader.pages, start=1):
            page_text = (page.extract_text() or "").strip()
            if page_text:
                parts.append(f"第 {index} 页\n{page_text}")
        return "\n\n".join(parts)
    except Exception as exc:
        raise HTTPException(400, f"PDF 解析失败：{exc}")


def knowledge_file_save_error_message(exc: SQLAlchemyError) -> str:
    detail = str(exc)
    if "Incorrect string value" in detail or "1366" in detail:
        return (
            "文件内容包含中文字符，但当前 MySQL 表或字段仍不是 utf8mb4。"
            "请重启后端让启动迁移生效；如仍失败，请执行："
            "ALTER DATABASE rag_system CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci; "
            "ALTER TABLE knowledge_files CONVERT TO CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci; "
            "ALTER TABLE knowledge_files MODIFY COLUMN content LONGTEXT CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;"
        )
    return "文件信息写入数据库失败，请稍后重试"


def chunk_text(text: str, file_id: int, chunk_size: int = 500, chunk_overlap: int = 50) -> list[dict]:
    """Split text into chunks with overlap and return them with ids."""
    if not text.strip():
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk_text_value = text[start:end].strip()
        if chunk_text_value:
            chunks.append({"id": f"{start}", "text": chunk_text_value})
        start += (chunk_size - chunk_overlap)
    return chunks

